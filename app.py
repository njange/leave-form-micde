from flask import Flask, render_template, request, send_file, url_for
from weasyprint import HTML
from docx import Document
import tempfile

app = Flask(__name__)

@app.route('/')
def home():
    return render_template('form-select.html')

@app.route('/form/<role>')
def form(role):
    # Render the appropriate form based on the role
    if role == 'staff':
        return render_template('form.html', role=role)
    elif role == 'hod':
        return render_template('form.html', role=role)  # Reuse the same form for HOD
    else:
        return "Invalid role", 404

@app.route('/generate/<role>', methods=['POST'])
def generate(role):
    data = request.form.to_dict()
    export_format = data.get("format", "pdf")

    # Determine the correct template based on the role
    if role == 'staff':
        template = 'letter_staff.html'
    elif role == 'hod':
        template = 'letter_hod.html'
    else:
        return "Invalid role", 404

    if export_format == "pdf":
        # Render the appropriate letter template
        rendered = render_template(template, data=data)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pdf_file:
            HTML(string=rendered, base_url=request.base_url).write_pdf(pdf_file.name)
            return send_file(pdf_file.name, as_attachment=True, download_name=f'{role}_leave_application.pdf')

    elif export_format == "docx":
        # Generate a DOCX file based on the role
        document = create_docx(data, role)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_file:
            document.save(docx_file.name)
            return send_file(docx_file.name, as_attachment=True, download_name=f'{role}_leave_application.docx')

def create_docx(data, role):
    doc = Document()
    doc.add_heading('REPUBLIC OF KENYA', 0)
    doc.add_paragraph('MINISTRY OF INFORMATION, COMMUNICATIONS AND THE DIGITAL ECONOMY')
    doc.add_paragraph('STATE DEPARTMENT FOR ICT AND DIGITAL ECONOMY')

    # Set the title based on the role
    if role == "hod":
        title = "LEAVE APPLICATION FORM FOR HEADS OF DEPARTMENTS"
    else:
        title = "LEAVE APPLICATION FORM FOR MEMBERS OF STAFF UNDER H.O.Ds"
    doc.add_heading(title, level=1)

    doc.add_paragraph("APPLICATION FOR ANNUAL LEAVE")
    doc.add_paragraph("(To be submitted at least 30 days before the leave is due to begin)\n")

    doc.add_paragraph(f"Name: {data['name']}")
    doc.add_paragraph(f"P/NO: {data['pno']}")
    doc.add_paragraph(f"Designation: {data['designation']}")
    doc.add_paragraph(f"I hereby apply for {data['days']} days annual leave beginning on {data['start_date']} to {data['end_date']}.")
    doc.add_paragraph(f"The last leave taken by me was from {data['last_from']} to {data['last_to']}.")
    doc.add_paragraph(f"Total leave days balance to date is {data['balance']} days.\n")

    doc.add_paragraph(f"While on leave, my contact will be: {data['contact']}, Tel: {data['phone']}")

    if data.get("payment") == "bank":
        doc.add_paragraph("Salary should continue to be paid into my bank account.")
    else:
        doc.add_paragraph("Salary should be paid at the following address:")
        doc.add_paragraph(data.get("payment_address", ""))

    doc.add_paragraph("I understand that I will require permission should I desire to spend leave outside Kenya in accordance to Human Resource policies and Procedures Manual 2016.")
    doc.add_paragraph("While on leave...")

    doc.add_paragraph("Date: ____________")
    doc.add_paragraph("Signature: ___________________________")

    doc.add_heading("PART II (To be completed by the Principal Secretary)", level=2)
    doc.add_paragraph("Approved / Not approved / Comments:")
    doc.add_paragraph("Date: ____________")
    doc.add_paragraph("Signed: ___________________________")
    doc.add_paragraph(f"{data['name']} will handle duties of my office.")
    doc.add_paragraph("PRINCIPAL SECRETARY")

    return doc

if __name__ == '__main__':
    app.run(debug=True)